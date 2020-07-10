#!/usr/bin/python3.6
# -*- coding: utf-8 -*-
# @Time       : 2020/7/10 16:35
# @Author     : 代登辉
# @Email      : 3276336032@qq.com
# @File       : TestDocX.py
# @Software   : PyCharm
# @Description: 方法描述,必填
import docx
from 在python中读取word文件 import word

docName = 'E:\课程\自然语言处理\代码\针对原始文本获取源数据和规范化\在python中读取word文件\sample-one-line.docx'

print('Document in full :\n', word.getTextWord(docName))
doc = docx.Document(docName)
print('Number of paragraphs :', len(doc.paragraphs))
print('Paragraph 2 :', doc.paragraphs[1].text)
print('paragraphs 2 style:', doc.paragraphs[1].style)

print('Paragraphs 1 :', doc.paragraphs[0].text)
print('Number of run in paragraphs 1 :', len(doc.paragraphs[0].runs))
for idx, run in enumerate(doc.paragraphs[0].runs):
    print('Run %s : %s' % (idx, run.text))

print('is Run 5 underlined :', doc.paragraphs[0].runs[5].underline)
print('is Run 1 bold :', doc.paragraphs[0].runs[1].bold)
print('is Run 3 italic :', doc.paragraphs[0].runs[3].italic)